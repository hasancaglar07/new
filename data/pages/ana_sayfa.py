# pages/ana_sayfa.py
import random
import os
import hashlib
import json
import streamlit as st
from dotenv import load_dotenv
from whoosh.index import create_in, open_dir
from whoosh.fields import Schema, TEXT, NUMERIC
from whoosh.qparser import QueryParser, MultifieldParser, AndGroup
from whoosh import index
from whoosh.analysis import StandardAnalyzer
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import sys
import locale
import openai
import uuid
import datetime
import traceback
import time
import fitz
from db import init_db, save_qa, get_all_qa
from utils import show_password_popup

# Encoding fix
sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')
locale.setlocale(locale.LC_ALL, 'tr_TR.UTF-8')
load_dotenv()

# --- UYGULAMA AYARLARI ---
st.set_page_config(page_title="Yediulya Kütüphanesi", page_icon="🕌", layout="wide")

# Bismillah ve Başlık
st.markdown('<div class="bismillah">بِسْمِ اللَّهِ الرَّحْمَٰنِ الرَّحِيمِ</div>', unsafe_allow_html=True)
st.markdown('<h1 class="main-title">🕌 Yediulya Kütüphanesi</h1>', unsafe_allow_html=True)

# --- YÜKLEME MESAJLARI VE DİNAMİK VERİLER ---
LOADING_MESSAGES = [
    "İlm-i Ledün madenini kazıyoruz... 📜",
    "Manevi detaylar derleniyor... 🌿",
    "Alıntılar tasnif ediliyor... 🕌",
    "İrfan sentezi hazırlanıyor... ✨",
    "Tasavvuf incileri toplanıyor... 💎",
    "Ruhani bilgiler yükleniyor... ☪️",
    "Hikmet kapıları açılıyor... 🔑"
]
TASAVVUF_TAVSIYELER = [
    "Kalbinizi rabıta ile arındırın.",
    "Zikr ile ruhunuzu yükseltin.",
    "Sabır, tasavvufun anahtarıdır.",
    "İlim, amelsiz faydasızdır.",
    "Mürşid, yol göstericidir.",
    "Tevekkül, huzurun kapısıdır.",
    "Şükür, nimetin artmasıdır."
]
HADISLER = [
    "Hadis: 'İlim öğrenmek her Müslümana farzdır.'",
    "Hadis: 'Allah'ı zikreden kalp diridir.'",
    "Hadis: 'Sabreden zafer bulur.'",
    "Hadis: 'En hayırlı amel, ihlastır.'"
]
AYETLER = [
    "Ayet: 'Allah sabredenlerle beraberdir.' (Bakara 153)",
    "Ayet: 'Zikrullah kalplerin huzurudur.' (Ra'd 28)",
    "Ayet: 'Rabbinizi zikredin.' (Araf 205)"
]
ORNEK_SORULAR = [
    "Rabıta nedir?",
    "Zikrullahın önemi?",
    "Mürşid-i Kamil özellikleri?",
    "Tasavvuf tarihi?",
    "Sufi yolları?",
    "Tevekkülün anlamı?",
    "İhlas nedir?"
]
os.makedirs("faiss_index", exist_ok=True)
os.makedirs("whoosh_index", exist_ok=True)
pdf_hash_file = "pdf_hash.json"
schema = Schema(
    book=TEXT(stored=True),
    author=TEXT(stored=True),
    page=NUMERIC(stored=True),
    content=TEXT(stored=True, analyzer=StandardAnalyzer()),
    pdf_file=TEXT(stored=True)
)
ix = create_in("whoosh_index", schema) if not index.exists_in("whoosh_index") else open_dir("whoosh_index")

# --- YARDIMCI FONKSİYONLAR ---
def get_pdf_hash(pdf_folder):
    hash_dict = {}
    for pdf_file in sorted(os.listdir(pdf_folder)):
        if pdf_file.endswith(".pdf"):
            file_path = os.path.join(pdf_folder, pdf_file)
            hasher = hashlib.md5()
            with open(file_path, "rb") as f:
                hasher.update(f.read())
            hash_dict[pdf_file] = hasher.hexdigest()
    return hash_dict

def load_saved_hash():
    if os.path.exists(pdf_hash_file):
        with open(pdf_hash_file, "r") as f:
            return json.load(f)
    return {}

def save_hash(hash_dict):
    with open(pdf_hash_file, "w") as f:
        json.dump(hash_dict, f)

@st.cache_resource(show_spinner=False, max_entries=10)
def build_data_havuzu():
    pdf_folder = "pdfler"
    current_hash = get_pdf_hash(pdf_folder)
    saved_hash = load_saved_hash()
    index_exists = (
        index.exists_in("whoosh_index") and 
        os.path.exists(os.path.join("faiss_index", "index.faiss")) and 
        os.path.exists(os.path.join("faiss_index", "index.pkl"))
    )
    rebuild = current_hash != saved_hash or not index_exists
    authors = set()
    docs = []
    
    def on_start():
        st.session_state.build_status = "started"
        st.session_state.build_message = random.choice(LOADING_MESSAGES)
        
    def on_progress(progress, pdf_file):
        st.session_state.build_status = "progress"
        st.session_state.build_progress = progress
        st.session_state.build_current_file = pdf_file
        
    def on_complete():
        st.session_state.build_status = "completed"
        st.session_state.build_message = "Veri havuzu başarıyla oluşturuldu!"
        st.session_state.build_progress = 1.0
        st.session_state.build_current_file = ""

    if rebuild:
        on_start()
        writer = ix.writer()
        from langchain_community.embeddings import HuggingFaceEmbeddings
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        from langchain.text_splitter import RecursiveCharacterTextSplitter
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
        from langchain_community.document_loaders import PyMuPDFLoader
        pdf_files = [f for f in os.listdir(pdf_folder) if f.endswith(".pdf")]
        
        for i, pdf_file in enumerate(pdf_files):
            progress = (i + 1) / len(pdf_files)
            on_progress(progress, pdf_file)
            
            base_name = pdf_file.replace(".pdf", "").replace("_", " ")
            if "-" in base_name:
                book_part, author_part = base_name.split("-", 1)
                book_name = book_part.strip().title()
                author = author_part.strip().title() + " Hz.leri"
            else:
                book_name = base_name.title()
                author = "Bilinmeyen Mürşid"
            authors.add(author.title())
            loader = PyMuPDFLoader(os.path.join(pdf_folder, pdf_file))
            pdf_docs = loader.load()
            for doc in pdf_docs:
                page_num = doc.metadata.get("page", 0) + 1
                content = doc.page_content.strip()
                if content:
                    writer.add_document(
                        book=book_name.lower(), 
                        author=author.lower(), 
                        page=page_num, 
                        content=content.lower(), 
                        pdf_file=pdf_file.lower()
                    )
                    doc.metadata["book"] = book_name
                    doc.metadata["author"] = author
                    doc.metadata["pdf_file"] = pdf_file
                    docs.append(doc)
        writer.commit()
        chunks = text_splitter.split_documents(docs)
        from langchain_community.vectorstores import FAISS
        vectorstore = FAISS.from_documents(chunks, embeddings)
        vectorstore.save_local("faiss_index")
        save_hash(current_hash)
        on_complete()
    else:
        from langchain_community.embeddings import HuggingFaceEmbeddings
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        from langchain_community.vectorstores import FAISS
        vectorstore = FAISS.load_local(
            "faiss_index", 
            embeddings, 
            allow_dangerous_deserialization=True
        )
        with ix.searcher() as searcher:
            all_docs = searcher.documents()
            for doc in all_docs:
                authors.add(doc.get("author").title())
                
    return vectorstore, sorted(list(authors))

def create_word_doc(data):
    if not data:
        return None
    doc = Document()
    doc.add_heading('Yediulya Kütüphanesi Sonuçları', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Tarih: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=len(data[0]))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, key in enumerate(data[0].keys()):
        hdr_cells[i].text = key
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(12)
    for row in data:
        row_cells = table.add_row().cells
        for i, value in enumerate(row.values()):
            row_cells[i].text = str(value)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(11)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def get_retriever(vectorstore, selected_authors):
    metadata_filter = {}
    if selected_authors:
        metadata_filter["author"] = {"$in": [a.lower() for a in selected_authors]}
    return vectorstore.as_retriever(
        search_kwargs={"k": 50, "filter": metadata_filter if metadata_filter else None}
    )

def highlight_query(text, query):
    if not query:
        return text
    for word in query.split():
        text = text.replace(word, f'<mark>{word}</mark>')
    return text

def show_page_image(pdf_path, page_num):
    try:
        doc = fitz.open(pdf_path)
        page = doc.load_page(page_num - 1)
        pix = page.get_pixmap(dpi=100)
        img_bytes = pix.tobytes("png")
        st.image(img_bytes, use_container_width=True)
        doc.close()
    except Exception as e:
        st.error(f"Sayfa resmi yüklenemedi: {str(e)}")

def display_assistant(message, query=None):
    result_type = message.get("result_type")
    if result_type == "Veri Arama":
        data = message.get("data", [])
        if data:
            st.subheader("📋 Tam Alıntı Listesi")
            if len(data) > 100:
                st.warning(f"📊 {len(data)} sonuç. İlk 100 gösteriliyor.")
                data = data[:100]
            st.info(f"🔍 {len(data)} adet alıntı bulundu")
            with st.expander("🔧 Sıralama & Görünüm", expanded=False):
                col1, col2 = st.columns([3, 1])
                with col1:
                    sort_option = st.selectbox(
                        "📊 Sıralama", 
                        ["Alaka Düzeyi", "Sayfa Numarası", "Kitap Adı"], 
                        key=f"sort_{message.get('unique_id')}"
                    )
                with col2:
                    view_option = st.radio("👁️ Görünüm", ["Kartlar", "Tablo"], horizontal=True, key=f"view_{message.get('unique_id')}")
            if sort_option == "Sayfa Numarası":
                data = sorted(data, key=lambda x: x["Sayfa"])
            elif sort_option == "Kitap Adı":
                data = sorted(data, key=lambda x: x["Kitap"])
            if view_option == "Kartlar":
                num_cols = 3
            else:
                num_cols = 1
            cols = st.columns(num_cols)
            for i, item in enumerate(data):
                with cols[i % num_cols]:
                    st.markdown(f"""
                    <div class="source-card">
                        <div class="card-header">
                            <h4>{item['Kitap']}</h4>
                            <span class="author-badge">{item['Yazar/Şahsiyet']}</span>
                        </div>
                        <div class="card-content">
                            <p class="page-info">📄 Sayfa: {item['Sayfa']}</p>
                            <p class="excerpt">{highlight_query(item["Tam Metin (Alıntı)"][:200] + "...", query)}</p>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    with st.expander("📖 Sayfa Resmini Gör"):
                        pdf_path = os.path.join("pdfler", item["PDF File"])
                        show_page_image(pdf_path, item["Sayfa"])
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            word_doc = create_word_doc(data)
            if word_doc:
                st.download_button("💾 Word İndir", word_doc, f"ilm_havuzu_{timestamp}.docx", key=message.get('unique_id'))
        else:
            st.info("📭 Sonuç bulunamadı.")
    elif result_type == "AI İrfan Sentezi":
        content = message.get("content", "")
        sources = message.get("sources", [])
        with st.chat_message("assistant"):
            container = st.empty()
            displayed_text = ""
            for char in content:
                displayed_text += char
                container.markdown(displayed_text)
        with st.expander("📚 Kaynaklar (Grid Kartlar)", expanded=False):
            if sources:
                books = list(set([s.get('book') for s in sources if s.get('book')]))
                authors = list(set([s.get('author') for s in sources if s.get('author')]))
                col1, col2, col3 = st.columns(3)
                col1.metric("📚 Kitap", len(books))
                col2.metric("👤 Yazar", len(authors))
                col3.metric("📎 Alıntı", len(sources))
            num_cols = 2
            cols = st.columns(num_cols)
            for i, doc in enumerate(sources):
                with cols[i % num_cols]:
                    st.markdown(f"""
                        <div class="source-card">
                            <div class="card-header">
                                <h4>Kaynak {i+1}</h4>
                                <span class="author-badge">{doc.get('author')}</span>
                            </div>
                            <div class="card-content">
                                <p class="book-info">📖 {doc.get('book')}</p>
                                <p class="page-info">📄 Sayfa: {doc.get('page')}</p>
                                <p class="excerpt">{highlight_query(doc.get('page_content')[:150] + "...", query)}</p>
                            </div>
                        </div>
                    """, unsafe_allow_html=True)
                    if 'pdf_file' in doc:
                        with st.expander("🔍 Sayfa Resmini Gör"):
                            pdf_path = os.path.join("pdfler", doc['pdf_file'])
                            show_page_image(pdf_path, doc.get('page'))
    else:
        st.markdown(message.get("content", ""))

def load_qa(question, answer):
    st.session_state.messages = [
        {"role": "user", "content": question},
        {"role": "assistant", "content": answer, "result_type": "AI İrfan Sentezi"}
    ]
    st.rerun()

def ask_data_havuzu(question: str, vectorstore, selected_authors, result_type):
    start_time = time.time()
    with st.spinner(random.choice(LOADING_MESSAGES)):
        try:
            retriever = get_retriever(vectorstore, selected_authors)
            assistant_message = {"role": "assistant", "result_type": result_type, "unique_id": str(uuid.uuid4())}
            if result_type == "Veri Arama":
                data = []
                with ix.searcher() as searcher:
                    parser = MultifieldParser(["author", "content"], schema=ix.schema, group=AndGroup)
                    query_parts = []
                    if selected_authors:
                        author_query = " OR ".join([f'author:"{a.lower()}"' for a in selected_authors])
                        query_parts.append(f"({author_query})")
                    if question:
                        query_parts.append(f'content:{question.lower()}')
                    full_query_str = " AND ".join(query_parts) if query_parts else "*"
                    q = parser.parse(full_query_str)
                    results = searcher.search(q, limit=100)
                    seen = set()
                    for hit in results:
                        unique_key = f"{hit['book']}_{hit['page']}_{hit['content'][:50]}"
                        if unique_key not in seen:
                            seen.add(unique_key)
                            data.append({
                                "Kitap": hit["book"].title(),
                                "Yazar/Şahsiyet": hit["author"].title(),
                                "Sayfa": hit["page"],
                                "Tam Metin (Alıntı)": hit["content"],
                                "PDF File": hit["pdf_file"]
                            })
                if question:
                    query_lower = question.lower()
                    data = sorted(data, key=lambda x: x["Tam Metin (Alıntı)"].lower().count(query_lower), reverse=True)
                assistant_message["data"] = data
            elif result_type == "AI İrfan Sentezi":
                deepseek_api_key = os.getenv("DEEPSEEK_API_KEY")
                if not deepseek_api_key:
                    raise ValueError("API anahtarı yok!")
                client = openai.OpenAI(base_url="https://api.deepseek.com", api_key=deepseek_api_key)
                relevant_docs = retriever.get_relevant_documents(question)
                context = "\n".join([doc.page_content for doc in relevant_docs])
                messages = [
                    {"role": "system", "content": """Sen bir tasavvuf âlimi ve uzmanısın. Cevabını SADECE verilen alıntılara dayalı oluştur; dış bilgi, genel tasavvuf bilgisi veya kişisel yorum ekleme. Verilen alıntılar dışındaki hiçbir üstat, kitap veya kavramı kullanma.
Sorulan kavramı tasavvuf bağlamında sentezle. Yanıtını şu yapıya göre organize et:
Anlam, Köken ve Önem: Kavramın anlamını, kökenini ve tasavvuftaki önemini kısaca sentezle (bu kısım yanıtın %20'sini geçmesin). Ardından, verilen alıntılardaki üstatlardan doğrudan alıntılarla destekle. Her alıntıda üstadın adını, kitabını ve sayfasını belirt (örneğin: "Ali Ramazan Dinç Efendi, Seyrsülük, sayfa X: 'Alıntı metni.'"). Eğer sayfa belirtilmemişse, sadece kitap adını belirt.
Görüşlerin Detaylandırılması ve Karşılaştırması: Verilen alıntılardaki üstatların görüşlerini detaylandır ve farklı bakış açılarını karşılaştır. Her detayı, ilgili üstatlardan doğrudan alıntılarla somutlaştır. Karşılaştırmalarda alıntıları yan yana koyarak analiz et (örneğin: "Hacı Hasan Efendi şöyle derken, Ali Ramazan Dinç Efendi şu şekilde farklı vurgular.").
Boyutların Analizi: Kavramın manevi, pratik ve felsefi boyutlarını analiz et. Her boyutu, verilen alıntılardaki üstat yazılarının doğrudan alıntıları ve örnekleriyle destekle. Analizini alıntılara dayandır, kendi sentezini minimum tut.
Yanıtının genel kuralları:
%80'i verilen alıntılardan DOĞRUDAN alıntılar olsun (tam metin alıntıları kullan, kısaltma).
%20'i sentez olsun (alıntıları bağlamak, karşılaştırmak veya özetlemek için).
Her paragrafta en az bir alıntı kullan ve üstadı alıntılayarak analiz et.
Yanıtı nesnel, akademik ve üstat odaklı tut; örnekler sadece verilen alıntılardan gelsin.
Eğer verilen alıntılarda yeterli detay yoksa, "Verilen eserlerde bu kavram için yeterli detay yok." diye belirt."""},
                    {"role": "user", "content": f"Alıntılar: {context}\nSoru: {question}"}
                ]
                stream = client.chat.completions.create(
                    model="deepseek-chat",
                    messages=messages,
                    temperature=0.5,
                    max_tokens=1500,
                    stream=True
                )
                answer = ""
                with st.chat_message("assistant"):
                    container = st.empty()
                    for chunk in stream:
                        if chunk.choices[0].delta.content is not None:
                            answer += chunk.choices[0].delta.content
                            container.markdown(answer)
                assistant_message["content"] = answer
                assistant_message["sources"] = [
                    {"book": doc.metadata.get("book"), "author": doc.metadata.get("author"), "page": doc.metadata.get("page"), "page_content": doc.page_content, "pdf_file": doc.metadata.get("pdf_file")}
                    for doc in relevant_docs if "pdf_file" in doc.metadata
                ]
                if answer:
                    save_qa(question, answer)
            st.session_state.messages.append(assistant_message)
            if len(st.session_state.messages) > 20:
                st.warning("⚠️ Geçmiş 20 mesajla sınırlı. Eski mesajlar silindi.")
                st.session_state.messages = st.session_state.messages[-20:]
            end_time = time.time()
            st.success(f"✅ Yanıt {end_time - start_time:.1f} saniyede hazırlandı")
            st.rerun()
        except Exception as e:
            st.error(f"❌ Hata oluştu: {str(e)}. Lütfen filtreleri kontrol edin veya tekrar deneyin.")
            with st.expander("🔧 Hata Detayı"):
                st.code(traceback.format_exc())
            if st.button("🔄 Tekrar Dene"):
                st.rerun()

# --- UYGULAMA AKIŞI ---
if "messages" not in st.session_state:
    st.session_state.messages = [{"role": "assistant", "content": "👋 Hoş geldiniz! Sorgu girin."}]
if not os.getenv("DEEPSEEK_API_KEY"):
    st.error("🔑 API anahtarı ekleyin!")
    st.stop()

if "build_status" not in st.session_state:
    st.session_state.build_status = "idle"
if "build_message" not in st.session_state:
    st.session_state.build_message = ""
if "build_progress" not in st.session_state:
    st.session_state.build_progress = 0.0
if "build_current_file" not in st.session_state:
    st.session_state.build_current_file = ""

with st.spinner(st.session_state.build_message or "Veri havuzu hazırlanıyor..."):
    progress_bar = st.progress(st.session_state.build_progress)
    if st.session_state.build_current_file:
        st.text(f"📖 İşleniyor: {st.session_state.build_current_file} ({int(st.session_state.build_progress * 100)}%)")
    
    vectorstore, authors = build_data_havuzu()
    
    if st.session_state.build_status == "completed":
        st.toast("✅ " + st.session_state.build_message, icon='🎉')
        st.session_state.build_status = "idle"
        st.session_state.build_message = ""
        st.session_state.build_progress = 0.0
        st.session_state.build_current_file = ""

st.session_state.vectorstore = vectorstore
st.session_state.authors = authors

st.markdown("---")
with st.expander("🔍 Filtreler", expanded=True):
    col1, col2, col3 = st.columns([3, 2, 2])
    with col1:
        selected_authors_inline = st.multiselect("Üstadlar", authors, key="ana_authors_select_inline", placeholder="Mürşid(ler) seçin...")
    with col2:
        result_type_inline = st.radio("🎯 Sonuç Türü", ["Veri Arama", "AI Sentezi"], key="ana_result_type_inline", horizontal=True)
    with col3:
        with st.popover("💡 Örnek Sorular"):
            st.markdown("**Hızlı Başlangıç:**")
            for soru in ORNEK_SORULAR:
                if st.button(soru, key=f"ornek_soru_{soru}"):
                    st.session_state.ana_chat_input_main = soru
                    st.rerun()

st.markdown("---")
with st.sidebar:
    st.markdown('<div class="sidebar-header">⚙️ Menü</div>', unsafe_allow_html=True)
    with st.expander("⭐ Favoriler", expanded=False):
        if "favorites" not in st.session_state:
            st.session_state.favorites = []
        if st.session_state.favorites:
            for i, fav in enumerate(st.session_state.favorites[-10:]):
                if st.button(f"⭐ {fav[:30]}...", key=f"fav_sidebar_{i}", use_container_width=True):
                    st.session_state.ana_chat_input_main = fav
                    st.rerun()
        else:
            st.info("📭 Favori yok")
        if st.button("📥 Sorguyu Favorile", key="add_to_favorites_sidebar", use_container_width=True):
            if "messages" in st.session_state and len(st.session_state.messages) > 1:
                current_query = st.session_state.messages[-2]["content"]
                if current_query and current_query not in st.session_state.favorites:
                    st.session_state.favorites.append(current_query)
                    st.toast("✅ Favorilere eklendi!", icon='✅')
    with st.expander("📚 Geçmiş Sorgular", expanded=False):
        qa_list = get_all_qa()[-20:]
        if qa_list:
            for i, qa in enumerate(qa_list):
                question, _, timestamp = qa
                unique_key = f"sidebar_qa_button_{i}_{timestamp}"
                if st.button(f"❓ {question[:35]}...", key=unique_key, use_container_width=True):
                    load_qa(question, qa[1])
        else:
            st.info("📭 Geçmiş yok")
    if st.button("🗑️ Geçmişi Temizle", key="clear_history_sidebar", use_container_width=True):
        st.session_state.messages = [{"role": "assistant", "content": "👋 Hoş geldiniz! Sorgu girin."}]
        st.rerun()
    with st.expander("❓ Bilgi & İpuçları", expanded=False):
        st.subheader("🕌")
        st.info(random.choice(TASAVVUF_TAVSIYELER))
        st.subheader("📜")
        st.info(random.choice(HADISLER))
        st.subheader("📖")
        st.info(random.choice(AYETLER))
        st.markdown("---")
        st.markdown("**İpuçları:**")
        st.caption("- Sorgunuzu yazdıktan sonra 'Enter' tuşuna basın.")
        st.caption("- Mobil cihazda yatay tutun.")
        st.caption("- API anahtarınızı kontrol edin.")

main_area = st.container()
with main_area:
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            if message["role"] == "user":
                st.markdown(message["content"])
            else:
                display_assistant(message, query=st.session_state.messages[-2]["content"] if len(st.session_state.messages) > 1 else None)
    if prompt := st.chat_input("💭 Sorgu girin (ör. rabıta)...", key="ana_chat_input_main"):
        st.session_state.messages.append({"role": "user", "content": prompt})
        ask_data_havuzu(prompt, vectorstore, selected_authors_inline, result_type_inline)

st.markdown('<div class="footer">© 2025 Yediulya Kütüphanesi</div>', unsafe_allow_html=True)